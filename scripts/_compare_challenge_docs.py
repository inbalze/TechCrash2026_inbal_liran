"""Compare challenge docs: .md vs .html vs .docx vs .pdf"""
import re, sys, hashlib
from pathlib import Path

ROOT = Path(r"c:\Projects\TechCrash2026\challenges")
MD  = ROOT / "challenges.md"
HTML= ROOT / "challenges_secret.html"
DOCX= ROOT / "challenges_secret.docx"
PDF = ROOT / "challenges_secret.pdf"

def from_md():
    return MD.read_text(encoding="utf-8")

def from_html():
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(HTML.read_text(encoding="utf-8"), "html.parser")
    # drop nav/style/script
    for t in soup(["style","script","nav"]): t.decompose()
    return soup.get_text("\n")

def from_docx():
    from docx import Document
    d = Document(DOCX)
    parts = []
    for p in d.paragraphs: parts.append(p.text)
    for tbl in d.tables:
        for row in tbl.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)

def from_pdf():
    from pypdf import PdfReader
    r = PdfReader(str(PDF))
    return "\n".join(p.extract_text() or "" for p in r.pages)

def norm(s):
    s = s.replace("\r","")
    s = re.sub(r"[ \t]+"," ", s)
    s = re.sub(r"\n{2,}","\n", s)
    return s.strip()

texts = {
    "md":   norm(from_md()),
    "html": norm(from_html()),
    "docx": norm(from_docx()),
    "pdf":  norm(from_pdf()),
}

# extract challenge titles + point values from each
def headings(t):
    # match e.g. "Challenge 3: Speed Loopback (50 + 200/150/100 pts)"
    rx = re.compile(r"Challenge\s+(\d+):\s*([^\n(]+?)(\(([^)]+)\))?\s*$", re.M)
    seen = []
    for m in rx.finditer(t):
        n   = m.group(1)
        ttl = m.group(2).strip()
        pts = (m.group(4) or "").strip()
        key = (n, ttl)
        if key not in [(a,b) for (a,b,_) in seen]:
            seen.append((n, ttl, pts))
    return seen

print("="*70)
for name, t in texts.items():
    print(f"\n--- {name.upper()} ({len(t)} chars) ---")
    for n,ttl,pts in headings(t):
        print(f"  Ch{n}: {ttl}  [{pts}]")

# scan each for known sensitive bits
print("\n" + "="*70)
print("Spot checks (does each include the same key sentences?)")
markers = [
    "SECRET",
    "DO NOT FORWARD",
    "all or nothing",
    "Performance + Time",
    "Completion + Time",
    "Volt-Meter",
    "Accelerometer 3D Cube",
    "FP8",
    "Frequency Detector",
    "PC Retro Game",
    "Press Right",
    "Speed Loopback",
    "FPGA Volt-Meter",
]
print(f"\n{'marker':40s} " + " ".join(f"{k:6s}" for k in texts))
print("-"*70)
for m in markers:
    row = [str(m.lower() in t.lower()) for t in texts.values()]
    print(f"{m:40s} " + " ".join(f"{r:6s}" for r in row))

# byte-content hashes (just so we can show they differ)
print("\nSHA256 of normalized text (different fingerprints expected):")
for k,t in texts.items():
    h = hashlib.sha256(t.encode("utf-8")).hexdigest()[:16]
    print(f"  {k:6s} {h}  len={len(t)}")
